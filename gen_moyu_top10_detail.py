# -*- coding: utf-8 -*-
"""gen_moyu_top10_detail.py — ST摸鱼风云-V2 · 深度跟读报告生成器
V168+G 九章式深度结构（对标中信目录规范） × V2评分刻度（十三维100分制，score_v2单源出分）
数据: st_scores_v2.json (2026-08-30 12:15 模型输出) + gen_moyu_top10_report.py 的 TOPS 单源(逐家logic/risk/dims)
命名: ST摸鱼风云-V2跟读报告_{rank:03d}_{code}_{简称}_{gen_date}.docx
结构: 投资要点页 + 九章（公司概况/财务分析/股权结构/司法信号/退市研判/壳价值/驱动压制/综合评级/附录）
"""
import ast, json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'C:\Users\xiaot\WorkBuddy\2026-05-16-task-2'

# ── 从 gen_moyu_top10_report.py 提取 TOPS（AST 安全解析，数据单源） ──
src = open(f'{BASE}\\gen_moyu_top10_report.py', encoding='utf-8').read()
tree = ast.parse(src)
TOPS = None
for node in tree.body:
    if isinstance(node, ast.Assign):
        for t in node.targets:
            if isinstance(t, ast.Name) and t.id == 'TOPS':
                TOPS = ast.literal_eval(node.value)
assert TOPS and len(TOPS) == 10, 'TOPS 提取失败'

v2 = json.load(open(f'{BASE}\\st_scores_v2.json', encoding='utf-8'))
meta, data = v2['meta'], v2['data']
by_code = {c['code']: c for c in data}

DIMS = [
    ('C1', '面值距离', 6), ('C2', '壳价值锚定', 8), ('S1', '实控人性质', 12), ('S2', '股权质押', 6),
    ('A1', '净资产', 10), ('A2', '扣非主营收入', 12), ('A3', '扣非净利润', 6), ('D1', '现金流', 4),
    ('B1', '立案/造假', 10), ('B2', '审计意见', 12), ('F2', '重组纾困', 6), ('F1', '财务趋势', 4), ('H1', '司法风险', 4),
]

# 五通道退市风险矩阵（V2十三维按退市通道归组；权重=5年176家退市案例实证贡献度）
CHANNELS = [
    ('市场通道', ['C1', 'C2']),
    ('财务红线', ['A1', 'A2', 'A3', 'D1']),
    ('规范通道', ['B2']),
    ('违法通道', ['B1']),
    ('股东/治理', ['S1', 'S2', 'H1']),
    ('纾困通道', ['F2', 'F1']),
]

# 每维档位解读（按得分区间给文本；档位键为「≤得分的最高档」）
TIER = {
    'C1': {6: '现价≥8元，距1元面值退市线极远，面值通道完全无威胁。',
           5: '现价在3~8元区间，距面值退市线安全边际较大，面值通道无威胁。',
           4: '现价在2~3元区间，距面值退市红线中等距离，需盯住低价波动。',
           3: '现价贴近面值退市红线（1~2元），面值退市是首要威胁维度之一。',
           0: '现价已跌破或濒临1元面值线，面值退市风险极端。'},
    'C2': {8: '市值≤壳基准(33.81亿)的5折，处于并购/借壳成本优势区——壳越便宜，重组方入场成本越低。',
           6: '市值处于壳基准5折~8折折价区，并购成本有吸引力。',
           4: '市值接近壳基准，平价区，壳价值中性。',
           2: '市值高于壳基准，并购溢价区，壳价值不突出。',
           0: '市值显著高于壳基准，壳价值无优势。'},
    'S1': {10: '国资实控（央企或省级国资委），保壳兜底能力与信用资源为最高档。',
           8: '市县国资实控，有地方国资信用兜底与资源协调能力。',
           3: '民企/个人实控，无国资兜底，保壳依赖自身资源与外部资本博弈。'},
    'S2': {6: '无质押/质押比例极低，控制权稳固，无平仓扰动。',
           3: '质押或冻结存在（历史口径缺失取中性档），控制权有潜在扰动。',
           0: '高质押或冻结，控制权不稳风险突出。'},
    'A1': {10: '归母净资产非常充裕（≥10分档），资不抵债风险极低。',
           8: '归母净资产充裕，无资不抵债风险，财务退市线安全。',
           6: '净资产尚可（正且有余量），安全边际一般。',
           4: '净资产偏薄，逼近资不抵债。',
           3: '净资产薄/资不抵债边缘，财务类退市风险显著。',
           0: '归母净资产为负（资不抵债），触及财务类退市红线。'},
    'A2': {12: '营收达标（主板≥3亿/双创≥1亿），财务类退市红线安全。',
           9: '营收接近达标但含金量打折（如低毛利冲量被模型识别）。',
           6: '营收低于红线但缺口可控。',
           3: '营收缺口大（触及「营收<红线+亏损」组合），已构成财务类退市风险。',
           0: '营收严重不达标（接近0或微收），财务类退市风险极端。'},
    'A3': {6: '扣非净利润为正，主业造血正常，盈利面无退市压力。',
           3: '微利或微亏，盈利面中性。',
           0: '扣非持续亏损，主业亏损面承压。'},
    'D1': {4: '经营现金流为正，造血正常。',
           2: '现金流接近平衡。',
           0: '经营现金流为负，资金链压力大。'},
    'B1': {10: '无立案/造假信号，重大违法退市通道安全。',
           7: '存在一般档行政处罚落地（无立案链条）或子公司被罚，非重大违法档。',
           0: '重大违法处罚落地（立案→处罚全链条），触发封顶（总分≤30）。'},
    'B2': {12: '标准无保留审计意见，审计面干净。',
           9: '带强调事项段的无保留意见，审计面基本干净但有提示事项。',
           6: '保留意见，审计瑕疵存在，需消除后才能摘帽。',
           0: '无法表示/否定意见，触发审计封顶。'},
    'F2': {6: '重整已进入执行/完成阶段（或重大资产重组落地），保壳确定性最高。',
           4: '重组/预重整推进中（预重整/重整/重大资产重组），保壳有实质抓手。',
           2: '有出售资产/债务豁免等保壳动作，但尚未形成重组级信号。',
           0: '无重组/纾困信号，保壳依赖自身经营。'},
    'F1': {4: '营收/扣非趋势改善，经营向上。',
           2: '趋势平稳，无显著恶化。',
           0: '趋势转弱/恶化，经营下行。'},
    'H1': {4: '实控人无冻结/限高/立案，司法面无风险。',
           2: '存在一定司法风险信号。',
           0: '实控人冻结/限高/立案，司法风险归零。'},
}

def tier_text(dk, score):
    m = TIER[dk]
    keys = sorted(m.keys(), reverse=True)
    for k in keys:
        if score >= k:
            return m[k]
    return m[min(m.keys())]

def classify_dims(dims):
    hi, mid, lo = [], [], []
    for (dk, dn, mx), sc in zip(DIMS, dims):
        tag = f'{dk} {dn}({sc}/{mx})'
        if sc >= mx * 0.8: hi.append(tag)
        elif sc <= mx * 0.5: lo.append(tag)
        else: mid.append(tag)
    return hi, mid, lo

def channel_matrix(dims):
    """五通道矩阵：返回 [(通道, 维度串, 得分合计, 满分, 状态)]，状态按得分占比≥80%安全/≤50%风险/中间关注"""
    dm = {dk: sc for (dk, dn, mx), sc in zip(DIMS, dims)}
    full = {dk: mx for dk, dn, mx in DIMS}
    out = []
    for ch, keys in CHANNELS:
        sub = [f'{dk}{dm[dk]}/{full[dk]}' for dk in keys]
        s = sum(dm[dk] for dk in keys)
        m = sum(full[dk] for dk in keys)
        ratio = s / m if m else 0
        st = '✅ 安全' if ratio >= 0.8 else ('⚠️ 风险' if ratio <= 0.5 else '👀 关注')
        out.append((ch, '、'.join(sub), s, m, st))
    return out

def h1(doc, txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

def h2(doc, txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(0x2e, 0x74, 0x9b)

def para(doc, txt, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.bold = bold
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p

def table(doc, rows, header=True):
    tb = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tb.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = tb.cell(i, j); c.text = ''
            r = c.paragraphs[0].add_run(str(cell))
            r.font.size = Pt(9)
            if i == 0 and header:
                r.bold = True
            elif i % 2 == 0:
                shd = c._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:fill'): 'EAF1F8'})
                c._tc.get_or_add_tcPr().append(shd)
    return tb

def warn_box(doc, lines):
    tb = doc.add_table(rows=1, cols=1)
    tb.style = 'Table Grid'
    c = tb.cell(0, 0)
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run('⚠️ ' + lines[0]); r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); r.font.size = Pt(10)
    for ln in lines[1:]:
        p2 = c.add_paragraph()
        r2 = p2.add_run(ln); r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0x7F, 0x5A, 0x1D)
    return tb

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = '（目录域：打开文档后 Ctrl+A → F9 刷新生成目录）'
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for e in (f1, it, f2, t, f3):
        run._r.append(e)

def risk_lines(risk, lo):
    """投资要点页/附录的核心风险条目（≤5条）"""
    lines = []
    if lo:
        lines.append('风险源维度：' + '、'.join(lo) + '，是主要失分区。')
    if '财务' in risk or '营收' in risk:
        lines.append('财务类退市红线临近（营收达标/扣非盈利/现金流任一维度失守）。')
    if 'B1=0' in risk or '封顶30' in risk:
        lines.append('重大违法处罚落地（B1=0），总分封顶30，退市锁定风险极高。')
    if 'B2=0' in risk or '封顶50' in risk:
        lines.append('审计无法表示/否定意见（B2=0），总分封顶50，规范风险重大。')
    if '面值' in risk:
        lines.append('股价贴近或跌破1元面值线，面值退市危机显著。')
    if '重组' in risk or '纾困' in risk:
        lines.append('暂无重组/纾困级正向信号，保壳依赖自身经营修复。')
    if len(lines) < 3:
        lines.append('需持续跟踪年报审计意见、立案进展与公告信号（近24个月巨潮定向采集）。')
    return lines[:5]

def build(code, rank, moyu, cheap, total, lv, mcap, price, br, cat, note, dims, logic, risk,
          out_dir=BASE, report_date_str='2026年9月1日', gen_date='20260901'):
    c = by_code[code]
    name = c['name']
    dm = {dk: (dn, mx, sc) for (dk, dn, mx), sc in zip(DIMS, dims)}
    hi, mid, lo = classify_dims(dims)
    ch_rows = channel_matrix(dims)
    LV_TXT = {
        'A': '保壳压力低、退市风险小，短期无强制退市通道命中的迹象。',
        'B': '保壳压力中等，个别维度存在失分，需跟踪下一报告期的修复情况。',
        'C': '保壳压力较大，存在明确风险敞口，若短板维度不修复，存在退市可能。',
        'D': '退市风险高，多通道濒临触发，除非出现实质性重整/纾困进展，否则应谨慎对待。'
    }

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 页眉页脚
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]; hp.text = ''
    hr = hp.add_run(f'ST摸鱼风云-V2 · 深度跟读报告｜{name}（{code}）'); hr.font.size = Pt(8); hr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fp = sec.footer.paragraphs[0]; fp.text = ''
    fr = fp.add_run('本报告由ST保壳评分系统V2模型生成（V168+G九章式结构 · V2评分刻度），仅供参考，不构成任何投资建议'); fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    # ── 封面 ──
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f'{name}\n{code}'); r.font.size = Pt(26); r.bold = True
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('ST摸鱼风云-V2 · 深度跟读报告'); r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run(f"V168+G 九章式深度结构 ｜ V2 评分刻度（榜单分数 = 报告分数）\n"
                   f"摸鱼榜第 {rank} 名 ｜ 摸鱼指数 {moyu:.1f} ｜ V2保壳分 {total} ｜ {lv}级 ｜ 全榜第 {br} 名\n"
                   f"报告日期：{report_date_str}\n"
                   f"榜单时点：{meta.get('updated', '')}（财务报告期 {meta.get('report_date', '')}）")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── 投资要点页（封面后、目录前，一页浓缩） ──
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{name}（{code}）投资要点'); r.font.size = Pt(18); r.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{c.get('board','')} ｜ {c.get('type','')} ｜ 报告期 {meta.get('report_date','')}"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    table(doc, [
        ['项目', '内容'],
        ['V2保壳分 / 等级', f'{total} 分 / {lv} 级（全榜第 {br} 名）'],
        ['摸鱼指数 / 排名', f'{moyu:.1f} / 第 {rank} 名（壳便宜分 {cheap}）'],
        ['市值 / 现价', f'{mcap:.2f} 亿元 / {price} 元'],
        ['实控人 / 属性', f'{c.get("controller","—")} / {c.get("controller_cat","—")}'],
        ['板块 / 类型', f"{c.get('board','—')} / {c.get('type','—')}"],
        ['报告期 / 数据时点', f"{meta.get('report_date','—')} / {meta.get('updated','—')}"],
    ])
    para(doc, '')
    p = doc.add_paragraph()
    r = p.add_run(f'综合评级：{lv}级（{total} / 100 分，V2十三维刻度）')
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) if lv in 'AB' else RGBColor(0xC0, 0x39, 0x2B)
    para(doc, '核心风险', bold=True, size=11)
    for i, ln in enumerate(risk_lines(risk, lo), 1):
        para(doc, f'{i}. {ln}', size=10)
    para(doc, '分析结论', bold=True, size=11)
    para(doc, f'{name}（{code}）当前V2保壳评分 {total} 分（{lv}级），全榜第 {br} 名。{LV_TXT[lv]}（详细研判见正文九章。）', size=10)
    para(doc, '分析机构：小调AI-WorkBuddy  |  分析日期：' + report_date_str, size=9, color=RGBColor(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── 目录 ──
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('目  录'); r.bold = True; r.font.size = Pt(14)
    add_toc(doc)
    doc.add_page_break()

    # ══ 第一章 公司概况 ══
    h1(doc, '一、公司概况')
    h2(doc, '1.1 基本信息')
    table(doc, [
        ['项目', '内容'],
        ['证券代码', code],
        ['证券简称', name],
        ['上市板块', c.get('board', '—')],
        ['ST类型', c.get('type', '—')],
        ['现价 / 市值', f"{price} 元 / {mcap:.2f} 亿元"],
        ['实控人 / 属性', f"{c.get('controller','—')} / {c.get('controller_cat','—')}"],
        ['全榜排名 / 等级', f"第 {br} 名 / {lv} 级（共 {len(data)} 家）"],
        ['报告期 / 榜单时点', f"{meta.get('report_date','—')} / {meta.get('updated','—')}"],
    ])
    h2(doc, '1.2 V2评分定位与退市通道速览')
    para(doc, f'本报告评分采用 ST保壳评分系统V2（十三维100分制，score_v2.py 单源出分，榜单分数=报告分数）。'
              f'{name} 当前 {total} 分、{lv} 级，全榜第 {br} 名。{LV_TXT[lv]}')
    table(doc, [['通道', '包含维度（得分/满分）', '通道得分', '通道满分', '状态']] + [list(x) for x in ch_rows])
    h2(doc, '1.3 报告口径说明')
    para(doc, f'· 财务数据：财务报告期 {meta.get("report_date", "")}；公告信号窗口：近24个月巨潮公告定向采集。\n'
              f'· 行情数据：榜单生成日（{meta.get("updated", "")[:10]}）收盘快照。\n'
              f'· 壳价值锚定：壳基准 33.81 亿（近2年154笔实控权变更市值中位数，own口径）。\n'
              f'· 营收红线：主板≥3亿 / 双创≥1亿（扣非口径）。\n'
              f'· 评级标准：A级>70｜B级51-70｜C级31-50｜D级≤30；联动规则：C1≤1时C2降档；涉造假立案S1封顶4。')

    # ══ 第二章 财务分析（V2口径） ══
    h1(doc, '二、财务分析（V2退市红线口径）')
    para(doc, '本章以V2财务四维（A1净资产/A2扣非主营收入/A3扣非净利润/D1现金流）替代V168的报表级财务分析——'
              'V2口径直接对标退市红线，比报表科目更贴近保壳判断。')
    table(doc, [
        ['维度', '得分', '满分', '解读'],
        ['A1 净资产', dm['A1'][2], 10, tier_text('A1', dm['A1'][2])],
        ['A2 扣非主营收入', dm['A2'][2], 12, tier_text('A2', dm['A2'][2])],
        ['A3 扣非净利润', dm['A3'][2], 6, tier_text('A3', dm['A3'][2])],
        ['D1 现金流', dm['D1'][2], 4, tier_text('D1', dm['D1'][2])],
    ])
    a2, a3 = dm['A2'][2], dm['A3'][2]
    if a2 <= 3 and a3 <= 3:
        verdict = '已触及「营收<红线+亏损」财务类退市红线组合——这是V2体系中最高优先级的财务退市信号，年报后触发*ST/退市风险高。'
    elif a2 <= 3:
        verdict = '扣非主营收入未达标（主板<3亿/双创<1亿），若叠加亏损将触发财务类退市红线，需重点跟踪年报。'
    elif a3 <= 3:
        verdict = '扣非持续亏损，盈利面承压，但营收达标使其未触及财务红线组合，核心矛盾在造血能力修复。'
    else:
        verdict = '财务四维均达标，未触及财务类退市红线，主业造血正常。'
    h2(doc, '2.1 财务红线组合判定')
    para(doc, verdict)
    h2(doc, '2.2 现金流与盈利质量')
    para(doc, f'D1 现金流 {dm["D1"][2]}/4：' + tier_text('D1', dm['D1'][2]))
    para(doc, f'A3 扣非净利润 {dm["A3"][2]}/6：' + tier_text('A3', dm['A3'][2]))

    # ══ 第三章 股权结构分析 ══
    h1(doc, '三、股权结构分析')
    h2(doc, '3.1 实控人属性（S1）')
    para(doc, f'实控人：{c.get("controller","—")}（{c.get("controller_cat","—")}）。S1 得分 {dm["S1"][2]}/12：' + tier_text('S1', dm['S1'][2]))
    h2(doc, '3.2 股权质押与冻结（S2）')
    para(doc, f'S2 得分 {dm["S2"][2]}/6：' + tier_text('S2', dm['S2'][2]))
    h2(doc, '3.3 控制权稳定性结论')
    s1, s2 = dm['S1'][2], dm['S2'][2]
    if s1 >= 10:
        concl = '国资实控且质押风险低，控制权稳定，具备保壳兜底能力。'
    elif s2 <= 3:
        concl = '质押/冻结存在，控制权有潜在扰动，需跟踪平仓风险与司法处置进展。'
    else:
        concl = '民企实控且无重大质押扰动，控制权平稳，但保壳无国资兜底。'
    para(doc, concl)

    # ══ 第四章 司法诉讼与公告信号 ══
    h1(doc, '四、司法诉讼与公告信号')
    h2(doc, '4.1 立案与重大违法（B1）')
    para(doc, f'B1 得分 {dm["B1"][2]}/10：' + tier_text('B1', dm['B1'][2]))
    h2(doc, '4.2 审计质量（B2）')
    para(doc, f'B2 得分 {dm["B2"][2]}/12：' + tier_text('B2', dm['B2'][2]))
    h2(doc, '4.3 实控人司法风险（H1）')
    para(doc, f'H1 得分 {dm["H1"][2]}/4：' + tier_text('H1', dm['H1'][2]))
    h2(doc, '4.4 公告信号面（近24个月巨潮定向采集）')
    if c.get('flags'):
        sigTxt = f'命中公告信号 {len(c["flags"])} 项，详见模型note：{note}'
    else:
        sigTxt = '近24个月无公告信号命中，监管与司法层面暂无新增事件。'
    para(doc, sigTxt)

    # ══ 第五章 退市风险研判与保壳路径 ══
    h1(doc, '五、退市风险研判与保壳路径')
    h2(doc, '5.1 五通道退市风险矩阵')
    para(doc, 'V2以退市通道为维度，权重源自5年176家强制退市案例实证贡献度（交易类55%/财务32%/规范8%/违法5-9%）。')
    table(doc, [['通道', '包含维度（得分/满分）', '通道得分', '通道满分', '状态']] + [list(x) for x in ch_rows])
    risky_ch = [x[0] for x in ch_rows if x[4].startswith('⚠️')]
    if risky_ch:
        para(doc, f'风险通道：{"、".join(risky_ch)}——是当前最可能触发退市的路径，需优先跟踪。')
    else:
        para(doc, '各通道均未进入风险区，短期无强制退市通道命中的迹象。')
    h2(doc, '5.2 保壳逻辑（模型解读）')
    para(doc, logic)
    h2(doc, '5.3 摘帽/保壳路径评估')
    para(doc, f'{name} 当前评级 {lv} 级。摘帽与否取决于短板维度修复：财务类看年报营收/扣非（A2/A3）、'
              f'规范类看审计意见（B2）与立案进展（B1）、市场类看面值（C1）。具体抓手见第六章壳价值与第七章驱动压制。')

    # ══ 第六章 资本运作成本与安全边际 ══
    h1(doc, '六、资本运作成本与安全边际')
    h2(doc, '6.1 壳价值锚定')
    para(doc, f'C2 壳价值得分 {dm["C2"][2]}/8：' + tier_text('C2', dm['C2'][2]))
    h2(doc, '6.2 并购成本与安全边际测算')
    if mcap > 0:
        para(doc, f'当前市值 {mcap:.2f} 亿元，对照壳基准 33.81 亿元（近2年154笔实控权变更市值中位数）：'
                  f'市值/基准 = {mcap/33.81*100:.0f}%。' +
                  ('处于并购/借壳成本优势区——壳越便宜，重组方入场成本越低，被并购概率上升。'
                   if mcap <= 33.81 * 0.5 else
                   '处于折价/平价区，并购成本有吸引力或中性。'
                   if mcap <= 33.81 else
                   '高于壳基准，壳价不占优，安全边际一般。'))
    else:
        para(doc, '市值数据缺失（未纳入摸鱼池），壳价值维度暂不评估。')
    h2(doc, '6.3 摸鱼视角')
    para(doc, f'摸鱼榜第 {rank} 名，摸鱼指数 {moyu:.1f}（壳便宜分 {cheap}）。摸鱼指数=50%×V2保壳分+50%×壳便宜分，'
              f'反映「保壳确定性 × 壳价吸引力」的综合博弈价值。')

    # ══ 第七章 核心驱动与压制因素 ══
    h1(doc, '七、核心驱动与压制因素')
    h2(doc, '7.1 保壳驱动因素')
    para(doc, '、'.join(hi) if hi else '（暂无满分档维度；保壳驱动主要依赖F2重组信号与国资兜底，当前均不强。）')
    h2(doc, '7.2 风险压制因素')
    para(doc, '、'.join(lo) if lo else '（无≤满分50%的风险源维度，整体风险可控。）')
    h2(doc, '7.3 多空总结')
    bull = hi[:2] or ['评级非D级（仍有保壳空间）']
    bear = lo[:2] or ['无重大风险源，但保壳确定性仍需公告验证']
    para(doc, f'多头逻辑：{"；".join(bull)}。空头逻辑：{"；".join(bear)}。综合评级见第八章。')

    # ══ 第八章 综合评级与落地策略 ══
    h1(doc, '八、综合评级与落地策略')
    h2(doc, '8.1 综合评级')
    table(doc, [
        ['项目', '结果'],
        ['V2保壳分 / 等级', f'{total} 分 / {lv} 级'],
        ['全榜排名 / 摸鱼排名', f'第 {br} 名 / 第 {rank} 名'],
        ['评级含义', LV_TXT[lv]],
    ])
    h2(doc, '8.2 分层策略建议')
    if lv == 'D':
        para(doc, '【持仓投资者】多通道濒临触发，除非出现实质性重整/纾困进展，否则应谨慎对待，跟踪退市风险提示公告。\n'
                  '【场外资金】不建议参与，退市风险与流动性风险偏高。\n'
                  '【产业方】若为壳价博弈，需确认财务红线与规范风险已出清。')
    elif lv == 'C':
        para(doc, '【持仓投资者】存在明确风险敞口，跟踪短板维度修复与年报窗口。\n'
                  '【场外资金】需等待实质性保壳动作（重组/重整/审计改善）确认后再评估。\n'
                  '【产业方】壳价进入可谈区间，但需先穿透财务红线与诉讼敞口。')
    else:
        para(doc, '【持仓投资者】保壳压力可控，跟踪下一报告期短板维度修复情况。\n'
                  '【场外资金】若参与壳价博弈，关注市值是否进入并购成本优势区（C2=8）。\n'
                  '【产业方】评级越高保壳确定性越强，壳价交易的安全边际越大。')
    h2(doc, '8.3 关键追踪时点')
    para(doc, f'· 年报窗口：下一报告期财务数据发布（营收/扣非/审计意见为核心变量）。\n'
              f'· 公告信号：近24个月巨潮定向采集窗口内的立案/处罚/重组/重整公告。\n'
              f'· 面值线：现价距1元面值退市线的距离变化（C1维度）。')

    # ══ 第九章 附录 ══
    h1(doc, '九、附录')
    h2(doc, '9.1 风险点清单')
    para(doc, risk, bold=True)
    h2(doc, '9.2 数据来源')
    para(doc, '巨潮资讯网（公告信号定向采集，近24个月）、东方财富（财务/股权数据）、腾讯财经（行情快照）、'
              '中登（质押周报）。财务报告期 ' + str(meta.get('report_date', '')) + '，榜单生成 ' + str(meta.get('updated', '')) + '。')
    h2(doc, '9.3 一页精简总结')
    para(doc, f'{name}（{code}）=V2评分 {total} 分/{lv}级/全榜第{br}名，摸鱼第{rank}名（指数{moyu:.1f}）。'
              f'{"、".join(hi) if hi else "无满分档维度"}为安全垫，'
              f'{"、".join(lo) if lo else "无极端风险源"}为风险源。'
              f'{LV_TXT[lv]}（详见正文九章。）')
    h2(doc, '9.4 风险提示与研究局限')
    warn_box(doc, [
        '本报告由模型基于公开数据自动生成，仅供研究参考，不构成任何投资建议。',
        '评分模型V2在点时回测（《ST保壳评分系统V2三年回测报告_20260830.docx》）中总分AUC≈0.53接近随机，'
        '判别力主要来自 F2重组信号(0.657) 与 C1面值距离(0.635)；B1/B2封顶机制存在选择效应，'
        '高保壳分≠一定摘帽、低保壳分≠一定退市，需结合年报原文与最新公告二次核验。',
        f'数据时点：榜单{meta.get("updated","")[:10]}生成（财务报告期{meta.get("report_date","")}），行情为当日快照；此后如有重大公告（立案/重组/审计更换）需重新评估。',
        '本报告为 V168+G 九章式深度结构 × V2 评分刻度：评分唯一口径为 ST保壳评分系统V2（十三维100分制），'
        '未使用 V168 六/八维评分，两套刻度语义不同，请勿与 V168 报告交叉比较。',
    ])
    para(doc, '')
    para(doc, '—— 报告完 ——', size=9)

    fn = f'{out_dir}\\ST摸鱼风云-V2跟读报告_{int(rank):03d}_{code}_{name.replace("*", "")}_{gen_date}.docx'
    doc.save(fn)
    return fn

if __name__ == '__main__':
    made = []
    for code, rank, moyu, cheap, total, lv, mcap, price, br, cat, note, dims, logic, risk in TOPS:
        fn = build(code, rank, moyu, cheap, total, lv, mcap, price, br, cat, note, dims, logic, risk)
        made.append(fn)
        print('已生成:', os.path.basename(fn))
    print(f'\n共 {len(made)} 份')
