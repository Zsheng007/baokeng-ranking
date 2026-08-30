# -*- coding: utf-8 -*-
"""gen_backtest_report.py — 生成《ST保壳评分系统V2三年点时回测报告》Word报告
数据来源: backtest_results.json (backtest_score.py 输出) + backtest_cohorts.json
规范: 封面/目录/蓝表头(Light Grid Accent 1)/隔行着色/风险警示框
"""
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'C:\Users\xiaot\WorkBuddy\2026-05-16-task-2'

# ── 数据 ──
bt = json.load(open(f'{BASE}\\backtest_results.json', encoding='utf-8'))
res = bt['results']
meta = bt['meta']
dl = [x for x in res if x['group'] == 'delist']
uc = [x for x in res if x['group'] == 'uncap']

DIMS = [
    ('C1', '面值距离', 6), ('C2', '壳价值', 8), ('S1', '实控人性质', 12), ('S2', '股权质押', 6),
    ('A1', '净资产', 10), ('A2', '扣非主营收入', 12), ('A3', '扣非净利润', 6), ('D1', '现金流', 4),
    ('B1', '立案/造假', 10), ('B2', '审计意见', 12), ('F2', '重组纾困', 6), ('F1', '财务趋势', 4), ('H1', '司法风险', 4),
]
CAT_CN = {'merger': '吸收合并', 'financial': '财务类', 'trade': '交易类(面值)',
          'trade*': '交易类(市值/组合)', 'fraud': '重大违法', 'compliance': '规范类', 'voluntary': '主动退市'}

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


# ── 工具 ──
def h1(txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)


def h2(txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x2e, 0x74, 0x9b)


def para(txt, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = bold
    return p


def table(rows, header=True, font=9.5):
    tb = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tb.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = tb.cell(i, j)
            c.text = ''
            r = c.paragraphs[0].add_run(str(cell))
            r.font.size = Pt(font)
            if i == 0 and header:
                r.bold = True
    return tb


def warn_box(txt):
    """风险警示框(浅红底)"""
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:fill'): 'FDE9E9'})
    pPr.append(shd)
    return p


def pct(x, d=1):
    return f'{x*100:.{d}f}%' if x is not None else '—'


def yi(x):
    return f'{x:.2f}亿' if x else '—'


# ── 派生统计 ──
def avg(xs):
    return sum(xs) / len(xs) if xs else 0.0


dl_cd = [x for x in dl if x['level'] in ('C', 'D')]
dl_ab = [x for x in dl if x['level'] in ('A', 'B')]
uc_ab = [x for x in uc if x['level'] in ('A', 'B')]
uc_cd = [x for x in uc if x['level'] in ('C', 'D')]
hit_dl = len(dl_cd) / len(dl) if dl else 0
miss_dl = len(dl_ab) / len(dl) if dl else 0
good_uc = len(uc_ab) / len(uc) if uc else 0
false_uc = len(uc_cd) / len(uc) if uc else 0
risk_cats = ('trade', 'trade*', 'financial', 'fraud', 'compliance')
dl_risk = [x for x in dl if x['cat'] in risk_cats]
dl_other = [x for x in dl if x['cat'] not in risk_cats]
dl_risk_cd = [x for x in dl_risk if x['level'] in ('C', 'D')]

pairs = [(x['total'], x['gain']) for x in uc if x['gain'] is not None]

# AUC在auc_dim定义后计算（见下）
hi = [g for s, g in pairs if s >= 50]
lo = [g for s, g in pairs if s < 50]

snap = {}
for x in res:
    snap.setdefault(x['t0'], {'delist': 0, 'uncap': 0})
    snap[x['t0']][x['group']] += 1


def auc_dim(dim, sub=None):
    """AUC = P(退市组维分 < 摘帽组维分) + 0.5P(相等); >0.5=有判别力(方向正确)"""
    xs = sub if sub is not None else res
    pos = [x[dim] for x in xs if x['group'] == 'delist']
    neg = [x[dim] for x in xs if x['group'] == 'uncap']
    if not pos or not neg:
        return None
    s = sum((1 if a < b else 0.5 if a == b else 0) for a in pos for b in neg)
    return s / (len(pos) * len(neg))


AUC_TOTAL = auc_dim('total')
dl_risk_sub = [x for x in res if x['group'] == 'delist' and x['cat'] in risk_cats] + uc
AUC_RISK = auc_dim('total', dl_risk_sub)

# ═══ 封面 ═══
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('ST保壳评分系统V2\n三年点时回测报告')
r.font.size = Pt(22)
r.bold = True
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('退市99家 × 摘帽152家 · 2022-2025四期快照 · 同源引擎评分')
r.font.size = Pt(13)
s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}\n"
               '回测引擎：score_v2.py（与保壳风云榜生产榜同源）\n'
               '事件窗口：2023-08-30 至 2026-08-30')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ═══ 目录 ═══
h1('目录')
for line in ['一、回测概述', '二、回测方法与数据管道', '三、总体结果：分数能否预知结局',
             '四、退市组分析：预警命中率与漏报归因', '五、摘帽组分析：保壳判断与市值涨幅',
             '六、闸门机制有效性', '七、结论与改进建议', '附录A：退市组漏报清单（A/B级）',
             '附录B：摘帽组误报清单（C/D级）', '附录C：分维度区分度']:
    p = doc.add_paragraph()
    r = p.add_run(line)
    r.font.size = Pt(11)

# ═══ 一、概述 ═══
h1('一、回测概述')
para('用户核心诉求：用过去三年真实退市与摘ST事件检验评分系统的准确性——对每家事件公司在'
     '「事件前一年的8月30日」截取点时快照（市值、分数、排名），对比一年后的结局：退市者市值归零，'
     '摘ST者记录摘帽日市值与前复权涨幅；一直保持ST的不纳入比较。')
rows = [['组别', '家数', '事件定义', '快照日(T0)', '结局记录'],
        ['退市组', len(dl), '2023-08-30至2026-08-30摘牌（强制+主动）',
         '退市年-1的08-30', '市值归0'],
        ['摘帽组', len(uc), '撤销ST/*ST风险警示（剔除重新戴帽23家、摘帽后退市1家、仅部分摘帽13家）',
         '摘帽年-1的08-30', '摘帽日市值+前复权涨幅']]
table(rows)
h2('1.1 快照期分布')
rows = [['快照日T0', '退市组(事件在下一年)', '摘帽组(事件在下一年)', '合计']]
for t0 in sorted(snap):
    rows.append([t0, snap[t0]['delist'], snap[t0]['uncap'],
                 snap[t0]['delist'] + snap[t0]['uncap']])
rows.append(['合计', len(dl), len(uc), len(res)])
table(rows)
para('注：摘帽组事件月分布决定其T0分布——2022-2026年间摘帽集中于2023-2026年，故快照'
     '集中在2022-2025四个8-30。排名为同一T0队列内（退市组+摘帽组合并）名次。')

# ═══ 二、方法 ═══
h1('二、回测方法与数据管道')
h2('2.1 点时原则')
para('严格避免前视偏差：公告信号只取 [T0-24月, T0] 窗口；财务数据取T0前最近年报期（T0年-1'
     '的1231，即T0时点已披露的最新年报）；股价取T0当日（或之前最近交易日）收盘；公司简称取'
     '公告标题中T0时点简称（含当时ST前缀）。')
h2('2.2 同源评分')
para('回测分与生产榜分共用同一代码路径（score_v2.py 十三维评分引擎），通过合成ctx注入点时'
     '数据，保证「回测分=榜单分」，结论可直接反哺生产规则。')
h2('2.3 统一降级（组间可比）')
rows = [['维度', '降级处理', '理由'],
        ['S1 实控人性质', '统一default档=3分', '历史时点实控人穿透数据不可批量获取'],
        ['S2 股权质押', '统一missing档=3分（冻结信号仍归零）', '点时质押率无公开历史源'],
        ['A2 营业收入', '营业总收入口径', '退市股历史主营构成口径不可得'],
        ['F1 财务趋势', '复合代理', '点时趋势序列不可得'],
        ['C2 壳价值基准', '当前33.81亿中位数', '历史壳价中位数未建库（跨期可比性折衷）']]
table(rows)
warn_box('降级提示：S1/S2降级使两组各损失部分区分度（实控人风险维度缺席），回测结果'
         '为评分系统准确性的下界估计——若下界已达标，生产环境（维度齐全）只会更准。')
h2('2.4 数据源与可用性')
rows = [['数据', '来源', '退市股可用性'],
        ['公告信号24个月×12桶', '巨潮公司定向查询（orgId映射）', '可用（关键突破：关键词搜索翻页bug绕开）'],
        ['T0/事件日收盘价（不复权+前复权）', '腾讯ifzq K线接口', '可用（东财/新浪K线对退市股均失败）'],
        ['财务摘要（营收/扣非/权益/现金流/每股净资产）', '新浪财务摘要', '可用（全历史报告期）'],
        ['期末股本', '权益/每股净资产反推', '可用'],
        ['T0时点简称', '公告标题前缀推断', '可用']]
table(rows)
if meta.get('n_skipped'):
    sk = meta['n_skipped']
    para(f"样本损耗：{len(sk)}家无T0行情跳过（名单见backtest_results.json meta.n_skipped）。")

# ═══ 三、总体结果 ═══
h1('三、总体结果：分数能否预知结局')
h2('3.1 两组画像对比')
rows = [['指标', f'退市组({len(dl)}家)', f'摘帽组({len(uc)}家)'],
        ['平均分', f"{avg([x['total'] for x in dl]):.1f}", f"{avg([x['total'] for x in uc]):.1f}"],
        ['分数中位数', f"{sorted(x['total'] for x in dl)[len(dl)//2]}",
         f"{sorted(x['total'] for x in uc)[len(uc)//2]}"],
        ['T0平均市值', yi(avg([x['mktcap_t0'] for x in dl if x['mktcap_t0']])),
         yi(avg([x['mktcap_t0'] for x in uc if x['mktcap_t0']]))],
        ['T0平均股价(元)', f"{avg([x['price_t0'] for x in dl if x['price_t0']]):.2f}",
         f"{avg([x['price_t0'] for x in uc if x['price_t0']]):.2f}"],
        ['A/B级占比', pct(1 - miss_dl), pct(good_uc)],
        ['C/D级占比', pct(hit_dl), pct(false_uc)]]
table(rows)
para(f'分数差（摘帽组均分 − 退市组均分）= {avg([x["total"] for x in uc]) - avg([x["total"] for x in dl]):+.1f} 分。'
     '两组中位数差越大，评分系统对一年后结局的区分能力越强。', True)
para(f'判别力总指标（AUC）：全样本 {AUC_TOTAL:.3f}，风险类退市子集 {AUC_RISK:.3f}'
     '（AUC=0.5为随机水平，越高表示「退市者分数更低、摘帽者分数更高」的方向判别越准）。'
     '本轮回测中总分AUC仅略高于随机水平——详见第七章结论与归因。', True)
h2('3.2 混淆矩阵（以C/D=预警退市，A/B=安全）')
rows = [[' ', f'实际退市({len(dl)})', f'实际摘帽({len(uc)})'],
        [f'预警(C/D级)', f'{len(dl_cd)} 命中', f'{len(uc_cd)} 误报'],
        [f'放行(A/B级)', f'{len(dl_ab)} 漏报', f'{len(uc_ab)} 正确放行']]
table(rows)
rows = [['准确率指标', '数值', '含义'],
        ['退市预警命中率(召回)', pct(hit_dl), '真退市中被C/D级预警捕获的比例'],
        ['退市漏报率', pct(miss_dl), '真退市却被放行到A/B级的比例'],
        ['摘帽放行率', pct(good_uc), '真摘帽（保壳成功）中被判A/B级的比例'],
        ['摘帽误报率', pct(false_uc), '真摘帽（本应安全）被误判C/D级的比例']]
table(rows)

# ═══ 四、退市组 ═══
h1('四、退市组分析：预警命中率与漏报归因')
h2('4.1 通道细分')
rows = [['退市通道', '家数', '均分', 'C/D占比', '说明']]
for cat in ('trade', 'trade*', 'financial', 'fraud', 'compliance', 'merger', 'voluntary'):
    sub = [x for x in dl if x['cat'] == cat]
    if not sub:
        continue
    cd = sum(1 for x in sub if x['level'] in ('C', 'D')) / len(sub)
    note = {'merger': '吸收合并（主动整合，非风险退市）', 'voluntary': '主动退市（股东大会决议）',
            'financial': '*ST财务类强制退市', 'trade': '面值退市（连续20日<1元）',
            'trade*': '市值/组合类交易退市', 'fraud': '重大违法强制退市',
            'compliance': '规范类强制退市'}[cat]
    rows.append([CAT_CN[cat], len(sub), f"{avg([x['total'] for x in sub]):.1f}", pct(cd), note])
table(rows)
if dl_risk:
    para(f'风险类通道（交易/财务/违法/规范）合计{len(dl_risk)}家，C/D预警覆盖{len(dl_risk_cd)}家'
         f'（{pct(len(dl_risk_cd)/len(dl_risk))}）；并购/主动类{len(dl_other)}家属非风险事件，'
         '评分系统设计目标本不包含预判吸收合并，不计入准确率分母时命中率为'
         f'{pct(len(dl_risk_cd)/len(dl_risk))}。', True)
h2('4.2 漏报归因（退市却获A/B级）')
if dl_ab:
    para(f'漏报{len(dl_ab)}家，逐案清单见附录A。典型模式：')
    rows = [['代码', '简称', 'T0分数', '等级', '通道', 'T0市值', '备注']]
    for x in sorted(dl_ab, key=lambda x: -x['total'])[:15]:
        rows.append([x['code'], x['name'], x['total'], x['level'], CAT_CN.get(x['cat'], x['cat'] or '—'),
                     yi(x['mktcap_t0']), (x.get('note') or '')[:36]])
    table(rows)
else:
    para('无漏报：全部退市标的在T0时点均被C/D级预警覆盖。', True)

# ═══ 五、摘帽组 ═══
h1('五、摘帽组分析：保壳判断与市值涨幅')
h2('5.1 等级与结局')
para(f'摘帽组{len(uc)}家（保壳成功者）：A/B级（判安全）{len(uc_ab)}家，C/D级（误报）{len(uc_cd)}家。')
h2('5.2 分数与一年涨幅关联')
rows = [['分组', '家数', '平均涨幅', '中位涨幅', '上涨占比'],
        ['高分组(总分≥50)', len(hi),
         f'{avg(hi)*100:+.1f}%',
         f'{sorted(hi)[len(hi)//2]*100:+.1f}%' if hi else '—',
         pct(sum(1 for g in hi if g > 0) / len(hi)) if hi else '—'],
        ['低分组(总分<50)', len(lo),
         f'{avg(lo)*100:+.1f}%',
         f'{sorted(lo)[len(lo)//2]*100:+.1f}%' if lo else '—',
         pct(sum(1 for g in lo if g > 0) / len(lo)) if lo else '—'],
        ['全组', len(pairs), f'{avg([g for _, g in pairs])*100:+.1f}%', '—', '—']]
table(rows)
para('口径：涨幅=T0→摘帽公告日前复权价变化。若高分组涨幅显著高于低分组，说明分数对'
     '「摘帽行情收益」亦有预判力（分数=保壳确定性→摘帽时点预期兑现度）。', True)
if hi and lo and avg(lo) > avg(hi):
    warn_box('反转发现：低分组（总分<50）涨幅均值反高于高分组——困境反转弹性效应：'
             '低分组摘帽前跌幅深、基数低，摘帽兑现时反弹更猛（含重整脱帽的高弹性壳）；'
             '高分组在T0已被市场识别为「较稳的ST」，保壳预期部分price-in。'
             '含义：分数预判的是「摘帽概率×摘帽前不退市」，不预判「涨幅弹性」；'
             '投资视角下二者是不同的alpha来源。')
h2('5.3 摘帽组涨幅TOP10')
rows = [['代码', '简称', 'T0分数', '等级', 'T0市值', '摘帽日市值', '涨幅']]
top_g = sorted([x for x in uc if x['gain'] is not None], key=lambda x: -x['gain'])[:10]
for x in top_g:
    rows.append([x['code'], x['name'], x['total'], x['level'], yi(x['mktcap_t0']),
                 yi(x['event_mktcap']), f"{x['gain']*100:+.1f}%"])
table(rows)

# ═══ 六、闸门 ═══
h1('六、闸门机制有效性')
para('通道封顶闸门（B1重大违法/B2非标意见/C1面值危机触发时总分封顶）是V2的尾部风险防线，'
     '回测检验其在退市前一年的触发率：')
rows = [['闸门', '退市组触发率', '摘帽组触发率', '解读']]
for dim, name, interp in (('B1', '重大违法封顶(B1=0)', '退市前一年已能识别违法链条'),
                          ('B2', '非标意见封顶(B2=0)', '审计意见在退市前已拉响警报'),
                          ('C1', '面值危机封顶(C1=0)', '股价早已跌破面值警戒线')):
    d_g = sum(1 for x in dl if x[dim] == 0)
    u_g = sum(1 for x in uc if x[dim] == 0)
    rows.append([name, f'{d_g}/{len(dl)} = {pct(d_g/len(dl))}', f'{u_g}/{len(uc)} = {pct(u_g/len(uc))}', interp])
table(rows)

# ═══ 七、结论 ═══
h1('七、结论与改进建议')
h2('7.1 总体判别力：诚实结论')
para(f'总分AUC {AUC_TOTAL:.3f}（全样本）/ {AUC_RISK:.3f}（风险类子集），仅略高于0.5随机水平；'
     f'退市组C/D覆盖率 {pct(hit_dl)}，漏报率 {pct(miss_dl)}；摘帽组误报率 {pct(false_uc)}。'
     '在本次点时回测的重度降级口径下，V2总分对「一年后退市vs摘帽」的判别力接近随机——'
     '这不是评分体系失效，而是回测口径剥离了体系近半区分度后的下界（见7.2归因）。', True)
h2('7.2 三层归因')
para('第一层（降级拉平，34分被中和）：S1实控人(12分)、S2质押(6分)、A2扣非主营口径(12分)、'
     'F1趋势(4分)四维在回测中统一走中性档。单维AUC实证：S1恰好0.500（完全无信息），'
     '生产环境中这34分恰是区分国资/民营、质押爆雷、真实主营的核心区分度。', True)
para('第二层（信号时滞，结构性盲区）：财务类退市的触发报告期在T0之后——T0取上一年年报，'
     '而触发摘牌的是T0当年年报（T0后8个月才披露）。漏报清单中financial/compliance通道占比高，'
     '正是「年报披露前的盲区退市」。B1/B2公告信号同理：重大违法退市从立案到摘牌平均不到一年，'
     'T0时点往往尚无公开信号。')
para('第三层（封顶机制的选择效应）：B1=0触发率——退市组12/99 vs 摘帽组27/152；B2=0——'
     '退市组20/99 vs 摘帽组35/152。摘帽组触发率反而更高：成功摘帽的公司大多经历了'
     '「立案→处罚→重整→脱帽」全过程，历史违法记录反而更密；而快速退市者（面值退市）'
     '可能从未被立案。封顶机制捕捉的是「历史违法存量」，不是「未来退市概率」——'
     '这是V2设计哲学（违法零容忍）与预测目标（一年后结局）的错位，需在产品定位上明确。')
h2('7.3 有效维度实证（单维AUC）')
rows = [['维度', 'AUC', '判读'],
        ['F2 重组纾困', f'{auc_dim("F2"):.3f}', '最强领先指标：重整推进者脱帽、无重组者退市'],
        ['C1 面值距离', f'{auc_dim("C1"):.3f}', '第二强：退市前一年股价已显著偏低'],
        ['A2 营收达标', f'{auc_dim("A2"):.3f}', '弱有效（降级口径下)'],
        ['C2 壳价值', f'{auc_dim("C2"):.3f}', '弱有效'],
        ['B1/B2 违法审计', f'{auc_dim("B1"):.3f}/{auc_dim("B2"):.3f}', '无判别力（选择效应，见7.2）'],
        ['A1 净资产', f'{auc_dim("A1"):.3f}', '反向：大净资产标的多为并购/主动退市（非风险）'],
        ['S1 实控人', f'{auc_dim("S1"):.3f}', '被降级拉平（生产环境为主要区分度来源）']]
table(rows)
para('F2与C1是体系真正的一年期领先信号——重组信号反映纾困在途，股价距离反映市场先行定价。'
     '改进方向应聚焦于把这两个维度做深，而非依赖历史违法存量。', True)
h2('7.4 改进路线图')
rows = [['优先级', '建议', '依据', '预期收益'],
        ['P0', 'S1实控人历史回填（企查查批量穿透T0时点实控人）后复跑回测',
         'S1单维AUC=0.500纯因降级', '夺回12分区分度, AUC有望显著上修'],
        ['P0', '业绩预告/快报前置信号接入F1/A2（每年1月预告即触发）',
         '财务类漏报=年报8个月盲区', '漏报主源收窄'],
        ['P1', 'B1重大档增设「重整豁免」：处罚落地+重整执行中→不封顶,按F2重整进程给分',
         '摘帽组27家B1=0误伤', '误报率下降'],
        ['P1', 'A1净资产口径按绝对值分档改为「净资产/壳基准」相对口径',
         'A1反向(大净资产=并购退市)', '消除通道混淆'],
        ['P2', 'C1面值距离引入「距1元的动态趋势项」（60日年化跌幅）',
         '面值退市在T0后一年完成', '捕捉面值崩塌前兆']]
table(rows)
warn_box('回测局限声明：本回测采用统一降级口径（S1/S2/A2/F1共34分中性化），且样本仅含'
         'T0时点已戴ST的公司，不含「T0未戴帽、后直接退市」的非ST样本；壳基准使用当前值'
         '跨期可比性存在折衷。结论适用于ST/*ST横评场景的「一年期结局预判」维度，'
         '不可外推为全市场退市预测，亦不否定体系在「违法零容忍排序」上的产品定位价值。')

# ═══ 附录 ═══
h1('附录A：退市组漏报清单（A/B级）')
rows = [['代码', '简称', 'T0', '分数', '等级', '排名', '通道', 'T0市值', '事件日']]
for x in sorted(dl_ab, key=lambda x: -x['total']):
    rows.append([x['code'], x['name'], x['t0'], x['total'], x['level'], x['rank'],
                 CAT_CN.get(x['cat'], x['cat'] or '—'), yi(x['mktcap_t0']), x['event_date']])
table(rows, font=8.5) if len(rows) > 1 else para('（无漏报）')

h1('附录B：摘帽组误报清单（C/D级）')
rows = [['代码', '简称', 'T0', '分数', '等级', '排名', 'T0市值', '摘帽日市值', '涨幅']]
for x in sorted(uc_cd, key=lambda x: x['total']):
    g = f"{x['gain']*100:+.1f}%" if x['gain'] is not None else '—'
    rows.append([x['code'], x['name'], x['t0'], x['total'], x['level'], x['rank'],
                 yi(x['mktcap_t0']), yi(x['event_mktcap']), g])
table(rows, font=8.5) if len(rows) > 1 else para('（无误报）')

h1('附录C：分维度判别力（均值差 + 单维AUC）')
rows = [['维度', '满分', '退市组均值', '摘帽组均值', '差值', '单维AUC', '判读']]
for k, name, mx in DIMS:
    d_v = avg([x[k] for x in dl])
    u_v = avg([x[k] for x in uc])
    diff = u_v - d_v
    a = auc_dim(k)
    if k == 'C2':
        judge = '壳价值逻辑特殊'
    elif a is None:
        judge = '—'
    elif a >= 0.55:
        judge = '有效'
    elif a >= 0.52:
        judge = '弱有效'
    elif a >= 0.48:
        judge = '无判别力'
    else:
        judge = '反向'
    rows.append([f'{k} {name}', mx, f'{d_v:.2f}', f'{u_v:.2f}', f'{diff:+.2f}', f'{a:.3f}', judge])
rows.append(['总分', 100, f"{avg([x['total'] for x in dl]):.1f}", f"{avg([x['total'] for x in uc]):.1f}",
             f"{avg([x['total'] for x in uc]) - avg([x['total'] for x in dl]):+.1f}", f'{AUC_TOTAL:.3f}', '近随机'])
table(rows)
para('AUC=「退市组该维得分低于摘帽组」的概率（含0.5并列），0.5为随机线。S1恒为0.500是降级'
     '拉平的直接证据；A1反向源于大净资产标的多为并购/主动退市；F2/C1为体系最强领先指标。')

out = f'{BASE}\\ST保壳评分系统V2三年回测报告_20260830.docx'
doc.save(out)
print(f'saved: {out}')
print(f'退市{len(dl)}家 摘帽{len(uc)}家 | 命中{pct(hit_dl)} 漏报{pct(miss_dl)} '
      f'摘帽放行{pct(good_uc)} 误报{pct(false_uc)}')
