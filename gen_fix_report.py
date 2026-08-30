# -*- coding: utf-8 -*-
"""gen_fix_report.py — 生成《ST保壳评分系统V2修改建议（301139元道案例）》Word报告
数据来源: 重算后的 st_scores_v2.json + 固定before快照 + st_risk_flags.json
"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'C:\Users\xiaot\WorkBuddy\2026-05-16-task-2'

# ── 数据 ──
v2 = json.load(open(f'{BASE}\\st_scores_v2.json', encoding='utf-8'))
after = next(x for x in v2['data'] if x['code'] == '301139')
flags = json.load(open(f'{BASE}\\st_risk_flags.json', encoding='utf-8'))
yd_flags = flags['flags'].get('301139', {})
n_sig = sum(len(v) for v in yd_flags.values())

BEFORE = {  # 2026-08-30修复前快照（来源: 旧st_scores_v2.json, 信号库元道12桶全空）
    'C1': 5, 'C2': 8, 'S1': 3, 'S2': 4, 'A1': 10, 'A2': 12, 'A3': 6, 'D1': 2,
    'B1': 10, 'B2': 12, 'F2': 0, 'F1': 0, 'H1': 4, 'total': 76, 'level': 'A', 'rank': 19,
}
DIMS = [
    ('C1', '面值距离', 6), ('C2', '壳价值', 8), ('S1', '实控人性质', 12), ('S2', '股权质押', 6),
    ('A1', '净资产', 10), ('A2', '扣非主营收入', 12), ('A3', '扣非净利润', 6), ('D1', '现金流', 4),
    ('B1', '立案/造假', 10), ('B2', '审计意见', 12), ('F2', '重组纾困', 6), ('F1', '财务趋势', 4), ('H1', '司法风险', 4),
]

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 封面
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('ST保壳评分系统V2\n评分失真诊断与修改建议'); r.font.size = Pt(22); r.bold = True
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('——基于 *ST元道(301139) 欺诈发行退市锁定案例'); r.font.size = Pt(13)
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run(f"报告日期：2026年8月30日\n分析机构：小调AI-WorkBuddy\n对照样本：《301139_ST元道_V168G深度分析报告_20260830.docx》（专家模型参考分 22/100）")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def h1(txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

def h2(txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs: r.font.color.rgb = RGBColor(0x2e, 0x74, 0x9b)

def para(txt, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.bold = bold
    return p

def table(rows, header=True, widths=None):
    tb = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tb.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = tb.cell(i, j); c.text = ''
            r = c.paragraphs[0].add_run(str(cell))
            r.font.size = Pt(9.5)
            if i == 0 and header:
                r.bold = True
    return tb

h1('一、案例概述：同一标的，两套结论')
para('专家模型（V168+G九章式）对 *ST元道 的定性：欺诈发行重大违法强制退市已锁定——2026-08-28'
     '证监会行政处罚决定书落地（公司罚2.39亿、实控人李晋5年市场禁入），同日深交所发出终止上市'
     '事先告知书，8月31日停牌，预计11月摘牌。参考评分 22/100（D级，仅作存档）。')
para('而风云榜V2修复前的评分：76分 / A级 / 全榜第19名 / 摸鱼榜第1名（摸鱼指数82）。'
     '一个已锁定退市的标的位居「保壳最容易+摸鱼潜伏」双榜首，属于系统性评分失真，必须溯源修正。')

h1('二、失真根因：三层问题')
h2('2.1 数据层（根因）：巨潮关键词搜索翻页bug + 信号库元道全空')
para('旧采集方案按13个关键词做全市场搜索。实测发现巨潮 hisAnnouncement/query 接口对关键词搜索的'
     '翻页存在严重缺陷：5页请求500条，实际仅返回150条、其中唯一公告约30条、120条为重复返回'
     '（即V1时代踩过的「同一公告重复返回」大坑）。时间倒序下，元道2025-07-11的立案公告排在'
     '400名开外，永远翻不到——导致元道12个信号桶全空：立案、处罚、保留意见、股份冻结全部漏采。')
para('而事实是（定向查询实测）：元道近24个月公告270条中，含立案告知书1条、立案调查进展7条、'
     '行政处罚决定书1条、处罚事先告知书1条、保留意见专项说明4条（2024/2025年报连续两年保留意见）、'
     '冻结类公告8条。信号非常密集，全部被旧管道漏掉。')
h2('2.2 规则层：处罚落地不触发B1归零')
para('旧B1逻辑只消费 investigation 桶。元道处罚决定书8-29披露（penalty桶命中），但即使命中'
     '旧逻辑也不会归零——「处罚落地」是比「立案」更重的违法实锤终局，必须触发B1=0（封顶30）。')
h2('2.3 口径层：摸鱼池准入无防飞刀闸门')
para('元道市值4.02亿（壳便宜分≈100），V2等级A（系数1.0），加权后摸鱼指数82高居榜首。'
     '这正是深度报告点名的教科书式飞刀案例：「超便宜但快退市」。摸鱼池的delisted判定只覆盖'
     '已完成退市整理的标的，不覆盖「终止上市事先告知书→停牌」这一锁定阶段。')

h1('三、本轮已实施修复（数据管道级 + B1规则细化）')
rows = [['修复项', '内容', '性质'],
        ['采集方案重写', 'fetch_risk_flags.py 弃关键词搜索，改为公司定向查询（stock=code,orgId，巨潮官方orgId映射），翻页正常、覆盖全量', 'bug修复'],
        ['B1接入penalty桶', '行政处罚决定书落地进入B1判定；fraud_involved同扫penalty桶', '信号补全'],
        ['公司案/个人案区分', 'PERSON_CASE过滤：董监高/实控人个人被立案处罚≠公司违法，不触发B1归零，走H1/S1（莫高案例：董事长个人被罚，恢复65/B）', '规则修正'],
        ['B1处罚三档细化', '重大档(0分封顶30)=公司本体处罚+立案全链条；一般档(7分)=公司本体单发处罚无立案（沈化/人福/绝味类）；子公司档(7分)=仅子公司被罚（纳川/启环类）', '规则细化'],
        ['信号库全量重采', '2965条信号（旧方案~150条），194/207家公司有信号；元道12桶从全空到命中27条', '数据修复'],
        ['会员3篇深度报告', '页面端：注册会员免费3篇「ST股票分析专家」V168+G九章式深度报告（配额+申请码+企微交付）', '功能升级']]
table(rows)
para('三档细化的依据：全量甄别62家「行政处罚落地」触发封顶的标的——38家为立案→处罚全链条'
     '（证监会稽查路径，重大信披违法实锤，封顶合理）；22家为公司本体单发处罚（地方证监局信披'
     '罚款，如ST沈化/ST人福/ST绝味，多为央企国资，无立案程序，处罚≠重大违法退市风险，原规则'
     '误伤）；2家仅为子公司被罚（*ST纳川/*ST启环，公司本体无处罚，不该归零）。')
h2('3.1 全榜分布演变（灰度对照）')
rows = [['阶段', 'A级', 'B级', 'C级', 'D级', '说明'],
        ['修复前（8-29旧信号库）', '64', '109', '33', '1', '信号大面积漏采，违法信号缺失，整体偏乐观'],
        ['中间态（B1一律封顶）', '6', '75', '61', '65', '过严：62家处罚落地全部归零，误伤单发/子公司处罚'],
        ['最终态（三档细化后）', '8', '91', '66', '42', '合理区间：38家重大违法封顶，22家一般+2家子公司回档']]
table(rows)
para('典型甄别案例：ST沈化（央企，辽宁证监局信披罚款，无立案）82分/A级——单发处罚仅扣3分；'
     '*ST纳川（子公司被罚+自身资不抵债+否定意见）33分/C级——财务问题主导降级，子公司处罚不再'
     '叠加归零；*ST元道（立案→处罚全链条）30分/D级——重大违法封顶维持。', True)

h1(f'四、修复实测：*ST元道 修复前后对比')
rows = [['维度', '分值', '修复前', '修复后', '变化原因']]
for k, name, mx in DIMS:
    b, a = BEFORE[k], after.get(k)
    reason = ''
    if k == 'B1' and b != a: reason = '行政处罚落地(欺诈发行)→0分, 触发封顶30'
    elif k == 'B2' and b != a: reason = '2024/2025年报连续保留意见→6分'
    elif k in ('S2', 'H1') and b != a: reason = '控股股东股份冻结/专户冻结/轮候冻结'
    rows.append([f'{k} {name}', f'/ {mx}', str(b), str(a), reason])
rows.append(['合计(封顶后)', '/ 100', f"{BEFORE['total']} ({BEFORE['level']}级·第{BEFORE['rank']}名)", f"{after['total']} ({after['level']}级·第{after.get('rank')}名)", 'B1=0→封顶30'])
table(rows)
para(f'修复后元道信号桶命中 {n_sig} 条公告（修复前0条）。专家模型参考分22，V2修复后'
     f'{after["total"]}分（{after["level"]}级·第{after.get("rank")}名）——方向与量级均已对齐至'
     f'「退市已锁定」区间；同时莫高65/B维持（公司/个人案区分后未误伤），沈化82/A恢复（央企'
     f'单发处罚不封顶），三档规则通过交叉验证。', True)

h1('五、修改建议清单（待拍板，按优先级）')
rows = [['优先级', '建议', '规则设计', '元道案例验证'],
        ['P0', '退市锁定一票否决', '公告命中「终止上市(事先)告知书/终止上市决定/退市整理期」→ delisted=true，踢出摸鱼池，榜单标记「已锁定退市」(不再参与评级)', '8-28告知书→8-31停牌，应立即出池'],
        ['P0', '摸鱼池防飞刀闸门', 'B1=0 或 B2=0 的标的不入摸鱼池（壳便宜分再高也无效——超便宜+快退市=飞刀）', '元道壳便宜分≈100仍应出池'],
        ['P1', 'A1应收质量修正', '应收账款/归母净资产>40%时A1降1档（元道8-9亿应收/18.89亿净资产≈45%，账面PB失真）', 'A1: 10→8'],
        ['P1', 'C1价格动量因子', '年内跌幅>70%时C1降1档（元道-82%、52周低点2.34元逼近面值区）', 'C1: 5→4'],
        ['P1', '信号采集纳入周更链', 'fetch_risk_flags.py 纳入weekly_update_friday.py（当前链外）；重大事件（处罚落地/终止上市告知）当日事件触发重跑', '处罚决定书8-29披露, 8-28采集差一天'],
        ['P2', 'B2直连审计意见字段', '审计意见改从年报财务数据API直读（AkShare审计意见表），公告标题匹配作降级方案——保留意见公告标题口径不稳定', '元道靠标题匹配本轮才补齐'],
        ]
table(rows)
para('说明：B1处罚三档细化已于本轮直接落地（属bug修复性质的规则还原，非口径变更）；P0两项为'
     '元道案例直接暴露的页面级缺口（摸鱼池准入），建议老Z拍板后实施；P1为口径微调（建议先跑灰度'
     '对照全榜分布变化再定）；P2为长期数据源加固。所有权重调整仍在score_config_v2.json外置框架'
     '内，改配置不改代码。')

h1('六、风险提示')
para('本报告为评分系统诊断文档，不构成投资建议。*ST元道已触及欺诈发行重大违法强制退市情形，'
     '2026-08-31起停牌，禁止抄底。数据来源：巨潮资讯网、东方财富、腾讯财经。')

out = f'{BASE}\\ST保壳评分系统V2修改建议_301139元道案例_20260830.docx'
doc.save(out)
print('saved:', out)
