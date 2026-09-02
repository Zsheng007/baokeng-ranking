# -*- coding: utf-8 -*-
"""生成《ST摸鱼风云-V5 报告模块 + 短信验证免费3篇 设计方案》Word 文档"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUT = r"ST摸鱼风云V5报告模块_短信验证3篇_设计方案.docx"

doc = Document()

# ============ 全局样式 ============
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)

GREEN = RGBColor(0x1A, 0x6B, 0x3A)
DARK = RGBColor(0x1A, 0x2B, 0x1F)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x88, 0x88, 0x88)

def shade_cell(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(str(text))
    r.font.size = Pt(size); r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = color

def add_h1(text):
    p = doc.add_heading('', level=1)
    r = p.add_run(text); r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = GREEN
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_h2(text):
    p = doc.add_heading('', level=2)
    r = p.add_run(text); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = DARK
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_h3(text):
    p = doc.add_heading('', level=3)
    r = p.add_run(text); r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLUE
    r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_para(text, bold=False, color=None, size=10.5, indent=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: r.bold = True
    if color: r.font.color.rgb = color
    if indent: p.paragraph_format.first_line_indent = Pt(21)
    return p

def add_bullet(text, color=None, bold=False, size=10):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color: r.font.color.rgb = color
    if bold: r.bold = True
    return p

def add_table(headers, rows, col_widths=None, header_fill='1A6B3A'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell(c, h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(c, header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    if col_widths:
        t.autofit = False
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    # 隔行着色
    for ri, row in enumerate(t.rows):
        if ri % 2 == 0 and ri > 0:
            for c in row.cells:
                shade_cell(c, 'F4FBF6')
    return t

def add_risk_box(title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    shade_cell(cell, 'FDF2F0')
    p0 = cell.paragraphs[0]
    r0 = p0.add_run('⚠ ' + title)
    r0.font.size = Pt(10.5); r0.bold = True; r0.font.color.rgb = RED
    for ln in lines:
        p = cell.add_paragraph()
        r = p.add_run(ln); r.font.size = Pt(9.5); r.font.color.rgb = DARK
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_bullet_table(title, rows):
    add_h3(title)
    add_table(['#', '项目', '说明', '负责人/依赖'], rows, col_widths=[1, 3.5, 6.5, 3.5])

# ============ 封面 ============
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ST摸鱼风云-V5 报告模块'); r.font.size = Pt(28); r.bold = True; r.font.color.rgb = GREEN
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('＋ 短信验证码 免费看3篇\n产品设计方案'); r.font.size = Pt(20); r.bold = True; r.font.color.rgb = DARK
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('—— 保壳风云榜（摸鱼榜）报告模块升级 · 2026-09-02 ——')
r.font.size = Pt(11); r.font.color.rgb = GREY
for _ in range(8): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('版本 V1.0 ｜ 编制：老Z ｜ 密级：内部'); r.font.size = Pt(10); r.font.color.rgb = GREY

doc.add_page_break()

# ============ 目录（手动） ============
add_h1('目录')
toc = [
    '1. 需求概述与产品目标',
    '2. 核心机制设计（短信验证 + 免费3篇）',
    '3. 短信服务商选型与成本测算',
    '4. 技术架构方案（前端 / 后端 / 数据管道）',
    '5. V5 报告模块内容设计',
    '6. 额度管理逻辑（3篇/手机号）',
    '7. 改造点清单（页面/脚本/部署）',
    '8. 实施路径（分阶段）',
    '9. 风险与合规提示',
    '10. 待决策事项',
]
for i, t in enumerate(toc, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}.  {t}'); r.font.size = Pt(12); r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i in (4, 5, 7, 8): r.font.color.rgb = GREEN

doc.add_page_break()

# ============ 1. 需求概述 ============
add_h1('1. 需求概述与产品目标')

add_h2('1.1 需求原文')
add_para('“除了评分榜单，我还需要将跟读报告模块内容换成ST摸鱼风云V5深度分析报告内容，只要对方输入手机号、短信验证，就可以免费看三篇文章，是否可以这样设计。”', color=DARK)
add_para('（原文经确认，“三篇文章”指同一手机号可免费查看的报告额度，即每家 ST 公司对应的 V5 深度报告。）')

add_h2('1.2 需求拆解')
add_table(
    ['#', '需求点', '含义', '状态'],
    [
        ['1', '保留评分榜单', '全名单/风云榜/摸鱼榜三视图 + 详情面板维持现状', '✅ 现状已满足'],
        ['2', '报告模块换内容', '从“跟读报告”换为「ST摸鱼风云-V5 深度分析报告」（九章式，V2评分刻度）', '🔄 本次升级'],
        ['3', '手机号 + 短信验证', '用户需输入11位手机号，通过短信验证码验证（不再是本地localStorage直接注册）', '🆕 本次新增'],
        ['4', '免费看3篇', '每个手机号累计可免费查看 3 份 V5 报告，超出需另付费或不可见', '🆕 本次新增'],
    ],
    col_widths=[1, 3.5, 8, 2.5]
)

add_h2('1.3 产品目标')
add_bullet('降低门槛：短信验证获取真实手机号，相比本地注册可防刷，沉淀真实客户联系方式')
add_bullet('内容升级：用 V5 九章式深度报告替代原有跟读报告，提升报告价值密度')
add_bullet('限量引流：“免费3篇”制造稀缺感，超额部分转化为付费/企微引流')
add_bullet('合规：短信验证需符合工信部实名制要求，收集手机号需有隐私政策声明')

doc.add_page_break()

# ============ 2. 核心机制设计 ============
add_h1('2. 核心机制设计（短信验证 + 免费3篇）')

add_h2('2.1 完整用户流程')
add_table(
    ['步骤', '用户操作', '系统行为', '触发点'],
    [
        ['①', '点击某公司报告按钮', '弹出“手机号验证”弹窗，要求输入11位手机号', '详情页「查看ST摸鱼风云-V5深度报告」'],
        ['②', '输入手机号，点“获取验证码”', '后端调短信服务商发验证码；校验限频（1条/分钟、5条/小时）', '点击获取验证码'],
        ['③', '输入6位验证码，点“验证并查看”', '后端校验验证码；校验该手机号剩余额度（已用<3则放行）', '点击验证'],
        ['④', '额度充足', '放行，页面打开对应 V5 报告（在线查看或弹出下载）', '展示报告'],
        ['⑤', '额度已用完', '提示“3篇免费额度已用完”，引导加企微客服付费解锁', '拦截并引导'],
        ['⑥', '累计', '服务端记录 {手机号: 已用篇数, 最近查看时间, 查看过的股票代码}', '每次放行后扣减'],
    ],
    col_widths=[1.2, 5, 6.3, 2.5]
)

add_risk_box('短信验证码安全要点', [
    '限频：同一手机号 1条/分钟、5条/小时、10条/天（阿里云硬限制，防止短信轰炸）',
    '有效期：验证码 5 分钟内有效，过期需重新获取',
    '防刷：可加图形验证码/滑块二次校验，防止自动化脚本批量刷码',
    '隐私：手机号仅用于报告解锁与额度记录，需在页面声明《隐私政策》，不得泄露',
])

add_h2('2.2 “免费3篇”额度规则')
add_bullet('额度单位：1 篇 = 1 只股票的 V5 深度报告')
add_bullet('周期：默认“永久累计3篇”（最简单）；也可做成“每月3篇”或“每注册用户3篇”')
add_bullet('存储：服务端按手机号维度记录（推荐），跨设备可同步；纯前端 localStorage 可被清除绕过（不推荐用于正式版）')
add_bullet('解锁：未用完时点开报告即扣减1篇（或“查看后扣减”，避免误点浪费）')

doc.add_page_break()

# ============ 3. 短信服务商选型 ============
add_h1('3. 短信服务商选型与成本测算')

add_h2('3.1 你关心的问题：需要支付费用吗？')
add_para('需要，但成本极低。短信验证码按条计费，量小时单价约 0.045-0.05 元/条。')
add_para('以“免费3篇”模型估算：每 1 个真实客户注册约消耗 1-3 条验证码（获取+可能重发），成本约 0.05-0.15 元/人。即使积累 1000 个客户，短信成本也仅约 50-150 元。', bold=True)

add_h2('3.2 主流短信服务商对比（验证码短信）')
add_table(
    ['服务商', '验证码单价', '个人可用', '接入方式', '备注'],
    [
        ['阿里云 SMS', '0.045元/条起', '⚠ 需企业实名', 'Dysmsapi SDK', '最主流，2026-05-20调价后0.045-0.05元/条'],
        ['腾讯云 SMS', '约0.05元/条', '⚠ 需企业实名', 'SMS SDK', '与腾讯系产品一致'],
        ['阿里云 短信认证', '按次', '✅ 个人可用', '直接回填验证', '个人自用推荐，免签名报备'],
        ['第三方聚合平台', '0.04-0.06元/条', '✅ 多数可个人', 'HTTP API', '如云片、Submail等'],
        ['极验/腾讯验证码', '按套餐', '✅ 个人', '前端SDK', '非短信，是图形/滑块验证，防刷用'],
    ],
    col_widths=[3, 2.8, 2.5, 3, 3.7]
)

add_risk_box('关键合规限制（重要）', [
    '工信部短信实名制：普通短信签名报备要求“企业/组织资质”，个人认证账号的自用资质难以通过签名实名制报备',
    '阿里云提示：“个人账号的自用资质无法通过签名实名制报备，个人用户请使用短信认证产品或升级为企业认证账号”',
    '结论：若以个人身份运营，推荐用「阿里云短信认证」产品（验证码自动回填，个人可用）；或注册企业/个体户账号以使用普通短信',
])

add_h2('3.3 成本测算（推荐方案）')
add_table(
    ['场景', '验证码条数/人', '单价', '单人成本', '1000人成本'],
    [
        ['基础（获取1次）', '1', '0.045元', '0.045元', '45元'],
        ['常规（含1次重发）', '2', '0.045元', '0.09元', '90元'],
        ['极端（多次重发）', '3', '0.045元', '0.135元', '135元'],
    ],
    col_widths=[4, 2.8, 2.2, 2.8, 2.8]
)
add_para('结论：短信成本可忽略不计，主要成本在“短信服务商认证门槛”而非费用本身。', bold=True, color=GREEN)

doc.add_page_break()

# ============ 4. 技术架构 ============
add_h1('4. 技术架构方案')

add_h2('4.1 架构总览')
add_table(
    ['层', '组件', '技术选型', '作用'],
    [
        ['前端', 'baokeng-rank.html / index.html', '纯静态HTML（现成）', '页面展示 + 报告按钮 + 验证弹窗'],
        ['后端', '短信验证 + 额度服务', 'FastAPI（Python）或云函数', '发验证码、校验、扣额度、记录手机号'],
        ['短信服务', '阿里云短信认证 / 腾讯云SMS', '第三方SDK', '发送验证码到手机'],
        ['数据管道', 'score_v2.py → st_scores_v2.json', '现有Python脚本', 'V2评分单源（榜单=报告）'],
        ['报告生成', 'st_moyu_fengyun_v5_report_template.py', 'python-docx模板', '生成V5九章式docx'],
        ['存储', '手机号额度表', 'SQLite/JSON文件/云数据库', '记录每手机号已用篇数'],
    ],
    col_widths=[2, 4, 3.5, 5.5]
)

add_h2('4.2 方案A：轻量 FastAPI 后端（推荐，自托管）')
add_bullet('复用现有 8001 端口（与“小调”FastAPI 一致）')
add_bullet('接口：POST /api/sms/send（发码）、POST /api/sms/verify（校验+扣额）、GET /api/quota/{phone}（查额度）')
add_bullet('额度存 SQLite 或 JSON 文件，简单可靠')
add_bullet('需部署到常驻服务器（本机/云主机），页面通过 fetch 调用')
add_bullet('优点：完全可控、免费部署；缺点：需常驻进程')

add_h2('4.3 方案B：Serverless 云函数（无服务器）')
add_bullet('腾讯云 SCF / 阿里云 FC，短信+额度逻辑写成云函数')
add_bullet('页面 HTTP 调用，无需常驻服务器')
add_bullet('优点：按调用付费、免运维；缺点：需学习云函数部署')

add_h2('4.4 方案C：纯前端模拟（最快落地，测试用）')
add_bullet('无真实短信，验证码固定为 123456（或前端随机生成后明文显示）')
add_bullet('额度存 localStorage，可被清除绕过')
add_bullet('优点：零后端、当天可跑通完整流程；缺点：不能用于正式对外（无法真实验证手机号、可绕过）')

add_risk_box('后端关键：为什么不能纯前端？', [
    '短信验证码必须由服务端发送（前端无法直接调短信服务商，密钥会泄露）',
    '额度扣减必须服务端记录，否则清除浏览器数据即可无限看',
    '真实手机号收集需有服务端存储与隐私保护',
])

doc.add_page_break()

# ============ 5. V5 报告模块 ============
add_h1('5. V5 报告模块内容设计')

add_h2('5.1 什么是 ST摸鱼风云-V5')
add_para('ST摸鱼风云-V5 是 ST摸鱼风云-V2 的定型升级版（2026-09-02 老Z定案），由「ST摸鱼风云-V2 跟读报告」升级定名而来。')
add_para('评分体系仍用 ST保壳评分系统V2（十三维100分制，score_v2 单源出分，与榜单同刻度）——所以榜单分数 = 报告分数不变。', bold=True)

add_h2('5.2 V5 九章式深度报告结构')
add_table(
    ['章', '标题', '内容要点'],
    [
        ['封面', '报告封面', '公司名/代码/板块/ST类型/戴帽日/评级'],
        ['投资要点', '一页浓缩', '核心数据表 + 评级 + 核心风险5条 + 结论'],
        ['第一章', '公司概况', '公司简介/主营业务/财务概况/股东背景'],
        ['第二章', '财务分析', '盈利能力/偿债/现金流/营运/趋势'],
        ['第三章', '股权结构', '前十大股东/实控人/股权质押'],
        ['第四章', '风险全景', '司法诉讼/戴帽原因/审计质量/退市风险研判'],
        ['第五章', '风险因素与投资逻辑', '驱动因素/压制因素/多空博弈'],
        ['第六章', '资本运作成本与安全边际', '隐性负债穿透/交易本质定性'],
        ['第七章', '估值分析', '壳资源估值/估值情景'],
        ['第八章', '投资建议与策略', 'V2十三维评分表/综合评级/分层策略/跟踪指标'],
        ['第九章', '附录', '风险跟踪表/财务汇总/股东结构/壳费参照/并购画像'],
        ['免责声明', '风险提示', '不构成投资建议，V2回测局限声明'],
    ],
    col_widths=[2.2, 4, 8.8]
)

add_h2('5.3 报告交付形态')
add_table(
    ['形态', '说明', '适用'],
    [
        ['docx 文件', 'V5模板生成，可下载/转发', '完整交付（推荐）'],
        ['在线 HTML', '页面内渲染阅读', '即时预览'],
        ['PDF', 'docx转PDF', '正式分发/打印'],
    ],
    col_widths=[2.5, 6, 6.5]
)

add_h2('5.4 报告库现状与升级计划')
add_table(
    ['项', '现状', '目标'],
    [
        ['全库数量', 'reports_baokeng_v2/ 207份（V2版，20260901）', '升级为 V5 九章式'],
        ['TOP10', '根目录10份 V2版', '本次先升 TOP10 为 V5'],
        ['命名', 'ST摸鱼风云-V2跟读报告_xxx', 'ST摸鱼风云-V5_简称_代码_报告.docx'],
    ],
    col_widths=[2.5, 7, 5.5]
)

doc.add_page_break()

# ============ 6. 额度管理 ============
add_h1('6. 额度管理逻辑（3篇/手机号）')

add_h2('6.1 数据模型')
add_table(
    ['字段', '类型', '说明'],
    [
        ['phone', 'string(11)', '用户手机号（主键）'],
        ['used_count', 'int', '已用篇数（0-3）'],
        ['last_view_at', 'datetime', '最近查看时间'],
        ['viewed_codes', 'list', '已查看的股票代码列表（防重复扣减/去重）'],
        ['created_at', 'datetime', '首次注册时间'],
        ['channel', 'string', '来源渠道（企微/网页等）'],
    ],
    col_widths=[3, 3, 9]
)

add_h2('6.2 扣减规则（推荐）')
add_bullet('用户点“验证并查看”→ 若 used_count < 3 → 放行，used_count+1，记录 viewed_codes')
add_bullet('同一股票重复查看：不重复扣减（viewed_codes 去重），避免误触浪费额度')
add_bullet('used_count ≥ 3 → 拦截，提示“3篇免费额度已用完”，引导加企微客服付费解锁（可接支付或线下）')
add_bullet('可选：每次查看生成一条申请记录（phone, code, 时间, 状态），便于客服跟进')

add_risk_box('防刷与风控', [
    '验证码限频 + 图形验证码双保险',
    '额度按手机号维度，无法通过清浏览器缓存绕过（服务端记录）',
    '异常高频（同一IP短时多次换号）可封IP',
    '收集手机号需符合《个人信息保护法》，页面声明用途',
])

doc.add_page_break()

# ============ 7. 改造点清单 ============
add_h1('7. 改造点清单')

add_h2('7.1 页面（generate_html.py → baokeng-rank.html/index.html）')
add_table(
    ['#', '改造点', '当前状态', '目标'],
    [
        ['1', '详情页报告按钮', '已删除（无报告按钮）', '重新加「查看ST摸鱼风云-V5深度报告」按钮'],
        ['2', '手机号验证弹窗', '已删除', '新增：手机号输入 + 获取验证码 + 验证'],
        ['3', '报告展示区', '已删除', '验证后展示/下载 V5 报告'],
        ['4', '额度提示', '无', '显示“免费3篇，已用X篇”'],
        ['5', '隐私声明', '无', '新增手机号使用声明'],
    ],
    col_widths=[1, 5, 4, 5]
)

add_h2('7.2 后端（新增）')
add_table(
    ['#', '模块', '说明'],
    [
        ['1', '短信发送接口', '调短信服务商，发验证码，限频'],
        ['2', '验证码校验接口', '校验验证码 + 查额度'],
        ['3', '额度扣减接口', '放行并 used_count+1'],
        ['4', '额度存储', 'SQLite/JSON/云数据库'],
        ['5', '企微对接预留', '超额时生成申请记录，客服企微跟进'],
    ],
    col_widths=[1, 4, 10]
)

add_h2('7.3 报告库（升级为V5）')
add_table(
    ['#', '改造点', '说明'],
    [
        ['1', 'TOP10 升级V5', '用 st_moyu_fengyun_v5_report_template.py 重跑TOP10'],
        ['2', '全库升级V5', '批量生成207份 V5 版（30秒内）'],
        ['3', '命名规范', 'ST摸鱼风云-V5_简称_代码_日期.docx'],
        ['4', 'score_v2 零偏差', '每份报告与 st_scores_v2.json 核对'],
    ],
    col_widths=[1, 4, 10]
)

add_h2('7.4 部署')
add_table(
    ['#', '项', '说明'],
    [
        ['1', '后端部署', 'FastAPI 8001 或 云函数'],
        ['2', 'GitHub Pages', '前端静态页照常托管'],
        ['3', 'CORS', '后端需允许页面域名跨域'],
        ['4', 'HTTPS', '短信接口需HTTPS'],
    ],
    col_widths=[1, 4, 10]
)

doc.add_page_break()

# ============ 8. 实施路径 ============
add_h1('8. 实施路径（分阶段）')

add_h2('阶段一：报告内容升级（先做，无后端依赖）')
add_table(
    ['步骤', '动作', '产出', '耗时'],
    [
        ['1', '用V5模板生成TOP10样例（冀凯等）', 'V5报告样例', '30分钟'],
        ['2', '全库207份升级V5', 'V5报告库', '30秒/批'],
        ['3', '确认命名与交付形态', '规范定稿', '评审'],
    ],
    col_widths=[1.5, 7, 4, 2.5]
)

add_h2('阶段二：页面加回报告模块（前端）')
add_table(
    ['步骤', '动作', '产出', '耗时'],
    [
        ['1', 'generate_html.py 加回报告按钮+弹窗', '页面可点报告', '1-2小时'],
        ['2', '短信验证弹窗 UI', '验证交互', '1小时'],
        ['3', '报告展示/下载', '交付V5报告', '1小时'],
    ],
    col_widths=[1.5, 7, 4, 2.5]
)

add_h2('阶段三：短信验证 + 额度后端（核心）')
add_table(
    ['步骤', '动作', '产出', '耗时'],
    [
        ['1', '选短信服务商（个人用短信认证/企业用普通短信）', '账号+密钥', '1-3工作日（认证）'],
        ['2', 'FastAPI 后端：发码/校验/扣额', '后端服务', '3-5小时'],
        ['3', '页面对接后端', '端到端跑通', '2小时'],
        ['4', '部署 + CORS + HTTPS', '上线', '半天'],
    ],
    col_widths=[1.5, 7, 4, 2.5]
)

add_h2('阶段四：验证与上线')
add_bullet('短信实测（真实手机号收码）')
add_bullet('3篇额度边界测试（第1/2/3/4篇）')
add_bullet('防刷测试（限频、重复验证码）')
add_bullet('GitHub Pages 部署 + 线上验证')

doc.add_page_break()

# ============ 9. 风险与合规 ============
add_h1('9. 风险与合规提示')
add_table(
    ['风险', '等级', '应对'],
    [
        ['短信实名制认证门槛（个人难报备签名）', '高', '用「短信认证」产品或注册企业/个体户账号'],
        ['短信成本（量小时极低）', '低', '0.045元/条，可控'],
        ['后端需常驻服务（方案A）', '中', '可用云函数免运维，或本机常驻'],
        ['个人信息合规（手机号收集）', '中', '页面声明隐私政策，最小化收集'],
        ['防刷/短信轰炸', '中', '验证码限频 + 图形验证'],
        ['3篇额度被绕过（纯前端）', '高', '必须服务端记录，勿用纯前端'],
    ],
    col_widths=[7, 2, 6]
)

add_risk_box('给老Z的一句话结论', [
    '短信验证码“免费看3篇”技术上完全可行，成本极低（千名客户约百元级），唯一前置门槛是短信服务商实名认证（个人建议用「短信认证」或注册个体/企业账号）。',
    '推荐路径：阶段一先升V5报告库（无后端、立即可做）→ 阶段二加回页面报告模块（前端）→ 阶段三接短信后端（认证+部署），分步上线、每步可验证。',
])

doc.add_page_break()

# ============ 10. 待决策 ============
add_h1('10. 待决策事项')
add_table(
    ['#', '决策点', '选项', '建议'],
    [
        ['1', '短信服务商', '阿里云短信认证/腾讯云SMS/聚合平台', '个人用阿里云短信认证'],
        ['2', '认证身份', '个人/企业/个体户', '视收款与报备需求'],
        ['3', '后端形态', 'FastAPI自托管/云函数/先纯前端模拟', '先纯前端跑通流程，再接真短信'],
        ['4', '报告升级范围', '仅TOP10/全库207份', '本次只升TOP10，后续可全库'],
        ['5', '免费篇数', '3篇/永久或每月', '默认永久累计3篇'],
        ['6', '超额变现', '企微付费/后续接支付', '先企微客服人工交付'],
    ],
    col_widths=[1, 3.5, 6, 4.5]
)

doc.add_page_break()

# ============ 附录 ============
add_h1('附录：短信验证核心接口设计')
add_table(
    ['接口', '方法', '入参', '出参'],
    [
        ['/api/sms/send', 'POST', '{phone}', '{code:0, msg:"发送成功"}'],
        ['/api/sms/verify', 'POST', '{phone, code}', '{code:0, msg:"验证成功", used:1, quota_left:2}'],
        ['/api/sms/check', 'GET', '?phone=', '{phone, used_count, quota_left, viewed_codes}'],
    ],
    col_widths=[3, 1.5, 5, 5.5]
)
add_para('说明：send 接口校验限频后调短信服务商；verify 校验验证码正确且未过期后，若 used_count<3 放行并 +1。')

# ============ 页脚 ============
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run('ST摸鱼风云-V5 报告模块 · 短信验证3篇 设计方案 ｜ 保壳风云榜 ｜ 2026-09-02')
    r.font.size = Pt(8); r.font.color.rgb = GREY

doc.save(OUT)
print('生成成功:', OUT)
print('大小:', os.path.getsize(OUT), 'B')
